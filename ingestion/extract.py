from __future__ import annotations

import logging
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

logger = logging.getLogger(__name__)


@dataclass(slots=True)
class ExtractedPage:
    page_number: int
    text: str
    source_kind: str = "text"


@dataclass(slots=True)
class ExtractedDocument:
    source_path: Path
    title: str
    extension: str
    pages: list[ExtractedPage]


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt", ".md"}

# A page with fewer than this many meaningful characters in its text layer is treated as an
# image-only page that needs OCR. Tunable via env so behaviour can be adjusted without code edits.
OCR_MIN_CHARS_PER_PAGE = int(os.getenv("OCR_MIN_CHARS_PER_PAGE", "20"))
# Hard cap on how many pages we will OCR per document, so a fully-scanned 700-page book cannot
# hang ingestion. Pages beyond the cap keep whatever (empty) text layer they had.
OCR_MAX_PAGES = int(os.getenv("OCR_MAX_PAGES", "100"))


def page_needs_ocr(text: str) -> bool:
    """A page needs OCR when its extracted text layer is empty/negligible (likely a scan/image)."""
    return len(text.strip()) < OCR_MIN_CHARS_PER_PAGE


def should_ocr_pdf(pages: list[ExtractedPage]) -> bool:
    """True if any page lacks a usable text layer. Kept for backwards compatibility."""
    return any(page_needs_ocr(page.text) for page in pages)


def extract_document(source_path: str | Path) -> ExtractedDocument:
    path = Path(source_path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported document type: {suffix or '<no extension>'}")

    handlers: dict[str, Callable[[Path], ExtractedDocument]] = {
        ".pdf": extract_pdf_document,
        ".docx": extract_docx_document,
        ".xlsx": extract_spreadsheet_document,
        ".xls": extract_spreadsheet_document,
        ".csv": extract_csv_document,
        ".txt": extract_text_document,
        ".md": extract_text_document,
    }
    return handlers[suffix](path)


def extract_text_document(path: Path) -> ExtractedDocument:
    return ExtractedDocument(
        source_path=path,
        title=path.stem.replace("_", " ").strip(),
        extension=path.suffix.lower(),
        pages=[ExtractedPage(page_number=1, text=path.read_text(encoding="utf-8"))],
    )


def extract_csv_document(path: Path) -> ExtractedDocument:
    import csv

    with path.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.reader(handle)
        lines = [", ".join(cell.strip() for cell in row) for row in reader]
    return ExtractedDocument(
        source_path=path,
        title=path.stem.replace("_", " ").strip(),
        extension=path.suffix.lower(),
        pages=[ExtractedPage(page_number=1, text="\n".join(lines))],
    )


def extract_docx_document(path: Path) -> ExtractedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                paragraphs.append(" | ".join(values))

    return ExtractedDocument(
        source_path=path,
        title=path.stem.replace("_", " ").strip(),
        extension=path.suffix.lower(),
        pages=[ExtractedPage(page_number=1, text="\n\n".join(paragraphs))],
    )


def extract_spreadsheet_document(path: Path) -> ExtractedDocument:
    import pandas as pd

    sheets = pd.read_excel(path, sheet_name=None)
    pages: list[ExtractedPage] = []

    for index, (sheet_name, frame) in enumerate(sheets.items(), start=1):
        cleaned = frame.fillna("")
        rows = []
        for _, row in cleaned.iterrows():
            values = [str(value).strip() for value in row.tolist() if str(value).strip()]
            if values:
                rows.append(" | ".join(values))
        page_text = f"Sheet: {sheet_name}\n" + "\n".join(rows)
        pages.append(ExtractedPage(page_number=index, text=page_text.strip()))

    return ExtractedDocument(
        source_path=path,
        title=path.stem.replace("_", " ").strip(),
        extension=path.suffix.lower(),
        pages=pages or [ExtractedPage(page_number=1, text="")],
    )


def extract_pdf_document(path: Path) -> ExtractedDocument:
    # 1) Extract the text layer per page (fast, no OCR).
    pages = extract_pdf_text_pages(path)

    # 2) OCR ONLY the pages whose text layer is empty/negligible (likely scans/images),
    #    capped at OCR_MAX_PAGES so a fully-scanned large document can't hang ingestion.
    ocr_targets = [p.page_number for p in pages if page_needs_ocr(p.text)]
    if ocr_targets:
        capped = ocr_targets[:OCR_MAX_PAGES]
        if len(ocr_targets) > OCR_MAX_PAGES:
            logger.warning(
                "Document %s: %d pages need OCR; OCR-ing first %d (OCR_MAX_PAGES).",
                path.name, len(ocr_targets), OCR_MAX_PAGES,
            )
        ocr_text = ocr_specific_pages(path, capped)
        by_number = {p.page_number: p for p in pages}
        for page_number, text in ocr_text.items():
            if text.strip():
                by_number[page_number] = ExtractedPage(
                    page_number=page_number, text=text.strip(), source_kind="ocr"
                )
        pages = [by_number[p.page_number] for p in pages]

    return ExtractedDocument(
        source_path=path,
        title=path.stem.replace("_", " ").strip(),
        extension=path.suffix.lower(),
        pages=pages,
    )


def extract_pdf_text_pages(path: Path) -> list[ExtractedPage]:
    import pdfplumber

    pages: list[ExtractedPage] = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            pages.append(ExtractedPage(page_number=index, text=(page.extract_text() or "").strip()))
    return pages


def ocr_specific_pages(path: Path, page_numbers: list[int]) -> dict[int, str]:
    """OCR only the given 1-based page numbers; returns {page_number: text}.

    Used by extract_pdf_document so only image-only pages are rendered+OCR'd, not the whole PDF.
    """
    import pypdfium2 as pdfium
    import pytesseract

    wanted = set(page_numbers)
    result: dict[int, str] = {}
    pdf = pdfium.PdfDocument(str(path))
    try:
        for index in range(len(pdf)):
            page_number = index + 1
            if page_number not in wanted:
                continue
            page = pdf[index]
            image = page.render(scale=2).to_pil()
            result[page_number] = pytesseract.image_to_string(image).strip()
    finally:
        pdf.close()
    return result


def ocr_pdf_pages(path: Path) -> list[ExtractedPage]:
    """OCR every page (full-document OCR). Retained for callers that want it explicitly."""
    import pypdfium2 as pdfium
    import pytesseract

    pdf = pdfium.PdfDocument(str(path))
    pages: list[ExtractedPage] = []
    for index in range(len(pdf)):
        page = pdf[index]
        image = page.render(scale=2).to_pil()
        text = pytesseract.image_to_string(image)
        pages.append(ExtractedPage(page_number=index + 1, text=text.strip(), source_kind="ocr"))
    return pages

